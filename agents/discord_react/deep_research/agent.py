
import discord
import logging
import asyncio
import json
import re
import os
import aiohttp
import urllib.parse
from bs4 import BeautifulSoup
from typing import Optional, List, Dict, Any, Tuple
from google.genai import types
from PIL import Image, ImageDraw

from agents.base_agent import BaseAgentSession

logger = logging.getLogger("DeepResearchAgent")


def generate_comparison_chart(fact_bank: List[Dict[str, Any]], filename: str = "chart.png") -> bool:
    logger.info("Initializing visual comparison graph compile pass...")
    
    latency_data = {}
    pricing_data = {}
    
    latency_pattern = re.compile(r'(?i)(apex|pebblehost|bisecthosting|bisect)\b.*?\b(\d+)\s*(?:-|to)?\s*(?:\d+)?\s*ms')
    price_pattern = re.compile(r'(?i)(apex|pebblehost|bisecthosting|bisect)\b.*?\b\$(\d+(?:\.\d+)?)')
    
    for item in fact_bank:
        fact = item.get("fact", "")
        
        lat_match = latency_pattern.search(fact)
        if lat_match:
            provider = lat_match.group(1).title()
            if "Bisect" in provider:
                provider = "Bisect"
            val = int(lat_match.group(2))
            if provider not in latency_data:
                latency_data[provider] = val
                
        pr_match = price_pattern.search(fact)
        if pr_match:
            provider = pr_match.group(1).title()
            if "Bisect" in provider:
                provider = "Bisect"
            val = float(pr_match.group(2))
            if provider not in pricing_data:
                pricing_data[provider] = val

    if not latency_data and not pricing_data:
        return False
        
    providers = list(set(list(latency_data.keys()) + list(pricing_data.keys())))
    if not providers:
        return False
        
    width, height = 800, 500
    img = Image.new("RGB", (width, height), "#1e1f22")
    draw = ImageDraw.Draw(img)
    
    draw.text((30, 20), "Minecraft Hosting Specs Comparison (US East)", fill="#ffffff")
    draw.line([(30, 45), (770, 45)], fill="#2f3136", width=2)
    
    y_start = 80
    row_height = 120
    
    for i, prov in enumerate(providers):
        y_pos = y_start + (i * row_height)
        
        draw.text((30, y_pos), prov, fill="#ffffff")
        
        lat = latency_data.get(prov, 0)
        if lat > 0:
            bar_width = min(300, int((lat / 100.0) * 300))
            draw.rectangle([(150, y_pos), (150 + bar_width, y_pos + 18)], fill="#5865f2")
            draw.text((155 + bar_width, y_pos + 3), f"{lat}ms Latency", fill="#949ba4")
            
        price = pricing_data.get(prov, 0.0)
        if price > 0:
            bar_width = min(300, int((price / 5.0) * 300))
            draw.rectangle([(150, y_pos + 25), (150 + bar_width, y_pos + 43)], fill="#248046")
            draw.text((155 + bar_width, y_pos + 28), f"${price:.2f}/mo IP", fill="#949ba4")
            
    try:
        img.save(filename)
        return True
    except Exception as img_err:
        logger.error(f"Failed to compile visual Pillow chart: {img_err}")
        return False



class DeepResearchSession(BaseAgentSession):
    def __init__(self, thread_id: int, user_id: int, prompt: str, loaded_contexts: str, channel: discord.Thread, depth: str = "standard", format_type: str = "markdown", flavor: str = "executive", domains_filter: str = "", guidelines: str = ""):
        super().__init__(thread_id, user_id, prompt, loaded_contexts, channel)
        self.depth = depth.lower().strip()
        self.format_type = format_type.lower().strip()
        self.flavor = flavor.lower().strip()
        self.domains_filter = domains_filter.strip()
        self.guidelines = guidelines.strip()
        
        self.sources: List[Dict[str, str]] = []
        self.fact_bank: List[Dict[str, Any]] = []
        self.scraped_urls: set = set()
        
        self.extract_sem = asyncio.Semaphore(3)
        
        self.plan_text: str = "Planning research targets..."
        self.current_thought: str = "Initiating active plan generation pass..."
        self.progress_msg: Optional[discord.Message] = None
        self.checklist_lines: List[str] = [
            "⚪ **Step 1: Planning research outline...**",
            "⚪ **Step 2: Crawling web indexes...**",
            "⚪ **Step 3: Scraping relevant pages & extracting facts...**",
            "⚪ **Step 4: Performing recursive gap analysis...**",
            "⚪ **Step 5: Compiling and formatting analytical report...**"
        ]

    def compile_react_transcript(self) -> str:
        return f"Deep Research Active. Depth: {self.depth}. Targets scraped: {len(self.scraped_urls)}."

    async def update_progress_msg(self, bot, active_idx: int, custom_header: str = ""):
        from agents.router_views import ViewResearchButton
        
        for i in range(len(self.checklist_lines)):
            line_clean = re.sub(r'^[⚪⏳✅]\s*', '', self.checklist_lines[i])
            if i < active_idx:
                self.checklist_lines[i] = f"✅ {line_clean}"
            elif i == active_idx:
                self.checklist_lines[i] = f"⏳ {line_clean}"
            else:
                self.checklist_lines[i] = f"⚪ {line_clean}"

        title_icon = "🔎" if active_idx < 4 else "📋"
        title_label = "Deep Research in Progress" if active_idx < 4 else "Research Concluding"
        
        card_content = (
            f"### {title_icon} {title_label}\n"
            f"----------------------------------------\n"
            f"{custom_header}\n\n"
            f"**Execution Checklist:**\n" + "\n".join(self.checklist_lines)
        )
        
        view = discord.ui.View(timeout=None)
        view.add_item(ViewResearchButton(self))
        
        if not self.progress_msg:
            self.progress_msg = await self.channel.send(content=card_content, view=view)
        else:
            try:
                await self.progress_msg.edit(content=card_content, view=view)
            except Exception as edit_err:
                logger.warning(f"Could not edit progress message: {edit_err}")

    async def execute_tick(self, bot):
        try:
            await self.run_research(bot)
        except Exception as e:
            logger.error(f"Deep Research Loop crashed: {e}", exc_info=True)
            self.status = "paused_error"
            await self.channel.send(f"🛑 **Research Interrupted:** An unexpected crash occurred during page analysis: `{e}`")

    async def run_research(self, bot):
        logger.info(f"Starting Deep Research loop in thread {self.thread_id}. Task: '{self.primary_task[:30]}'")
        
        await self.update_progress_msg(bot, active_idx=0, custom_header="*Formulating deconstruction plan and targeting sub-questions...*")
        queries = await self.generate_queries(bot)
        
        self.plan_text = f"Drafted {len(queries)} initial query targets across public indexes:\n" + "\n".join([f"- {q}" for q in queries])
        self.current_thought = "Executing parallel search queries..."
        
        await self.update_progress_msg(bot, active_idx=1, custom_header=f"📡 *Crawling web indexes concurrently using `{len(queries)}` target queries...*")
        
        search_tasks = [self.search_duckduckgo(query) for query in queries]
        search_results = await asyncio.gather(*search_tasks)
        
        candidates = []
        seen_links = set()
        
        exclude_domains = []
        if self.domains_filter:
            exclude_domains = [d.replace("-", "").strip().lower() for d in self.domains_filter.split(",") if d.strip().startswith("-")]

        for res_list in search_results:
            for item in res_list:
                url = item.get("url", "")
                if url in seen_links or url in self.scraped_urls:
                    continue
                    
                domain_lower = url.lower()
                if any(ex in domain_lower for ex in exclude_domains if ex):
                    continue
                    
                seen_links.add(url)
                candidates.append(item)
                
        if not candidates:
            candidates = await self.search_duckduckgo(self.primary_task)

        await self.update_progress_msg(bot, active_idx=2, custom_header=f"📄 *Relevance filtering candidates. Selecting top sources...*")
        
        max_links = 5 if self.depth == "brief" else 15 if self.depth == "standard" else 30 if self.depth == "exhaustive" else 60
        selected_candidates = await self.filter_relevance(bot, candidates[:40], limit=max_links)
        
        self.current_thought = f"Selected {len(selected_candidates)} highly relevant pages to scrape. Extracting atomic facts..."
        await self.update_progress_msg(bot, active_idx=2, custom_header=f"📄 *Concurrently scraping `{len(selected_candidates)}` pages and analyzing contents...*")
        
        scrape_tasks = [self.process_single_page(bot, cand.get("url"), cand.get("title", "")) for cand in selected_candidates]
        await asyncio.gather(*scrape_tasks)
        
        if self.depth != "brief" and len(self.sources) > 0:
            self.current_thought = "Analyzing database for knowledge gaps and discrepancies..."
            await self.update_progress_msg(bot, active_idx=3, custom_header="🔍 *Performing analytical gap verification on facts database...*")
            
            gap_queries = await self.analyze_gaps(bot)
            if gap_queries:
                self.current_thought = f"Identified gaps. Executing recursive query wave: {gap_queries}"
                await self.update_progress_msg(bot, active_idx=3, custom_header=f"📡 *Recursive pass: Executing `{len(gap_queries)}` targeted lookup queries...*")
                
                gap_search_tasks = [self.search_duckduckgo(g_q) for g_q in gap_queries]
                gap_search_results = await asyncio.gather(*gap_search_tasks)
                
                gap_candidates = []
                for res_list in gap_search_results:
                    for item in res_list:
                        url = item.get("url", "")
                        if url not in self.scraped_urls and url not in [c.get("url") for c in gap_candidates]:
                            gap_candidates.append(item)
                            
                if gap_candidates:
                    await self.update_progress_msg(bot, active_idx=3, custom_header=f"📄 *Scraping secondary gap-filling pages...*")
                    gap_scrape_tasks = [self.process_single_page(bot, cand.get("url"), cand.get("title", "")) for cand in gap_candidates[:4]]
                    await asyncio.gather(*gap_scrape_tasks)
                    
        self.current_thought = "Synthesizing all gathered notes into final master report..."
        await self.update_progress_msg(bot, active_idx=4, custom_header="✍️ *Synthesizing collected fact indexes and generating Markdown document...*")
        
        report_markdown = await self.compile_report_sequential(bot)
        
        chart_generated = generate_comparison_chart(self.fact_bank, "chart.png")
        
        slug = re.sub(r'[^a-zA-Z0-9]', '_', self.primary_task[:30]).strip("_")
        
        if self.format_type == "docx":
            filename = f"Research_Report_{slug}.docx"
            try:
                import docx
                doc = docx.Document()
                
                doc.add_heading(f"Research Report: {self.primary_task}", level=0)
                
                if chart_generated:
                    doc.add_heading("Visual Comparison Data Matrix", level=1)
                    doc.add_picture("chart.png", width=docx.shared.Inches(5.5))
                    doc.add_paragraph("Figure 1.0: Dynamically compiled metrics chart.")
                
                sections = report_markdown.split("### ")
                for sec in sections:
                    if not sec.strip():
                        continue
                    lines = sec.strip().split("\n")
                    heading = lines[0]
                    doc.add_heading(heading, level=2)
                    
                    body_text = "\n".join(lines[1:])
                    doc.add_paragraph(body_text.strip())
                    
                doc.save(filename)
                file = discord.File(fp=filename, filename=filename)
                
                files_to_send = [file]
                if chart_generated:
                    files_to_send.append(discord.File(fp="chart.png", filename="chart.png"))
                    
                mention_prefix = f"🛎️ <@{self.user_id}>, " if self.depth == "extreme" else ""
                
                await self.channel.send(
                    content=f"{mention_prefix}### 📋 Deep Research Summary Report Completed!\n"
                            f"I have successfully searched and compiled data from **{len(self.sources)} websites** "
                            f"into a polished Word document (.docx). The dossier has been attached below.",
                    files=files_to_send
                )
                
                if os.path.exists(filename):
                    os.remove(filename)
                if os.path.exists("chart.png"):
                    os.remove("chart.png")
                    
            except Exception as docx_err:
                logger.error(f"Failed compiling Word Document: {docx_err}. Falling back to standard Markdown compilation...")
                self.format_type = "markdown"
                
        if self.format_type == "markdown":
            filename = f"Research_Report_{slug}.md"
            try:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(report_markdown)
                    
                file = discord.File(fp=filename, filename=filename)
                files_to_send = [file]
                if chart_generated:
                    files_to_send.append(discord.File(fp="chart.png", filename="chart.png"))
                    
                mention_prefix = f"🛎️ <@{self.user_id}>, " if self.depth == "extreme" else ""
                
                await self.channel.send(
                    content=f"{mention_prefix}### 📋 Deep Research Summary Report Completed!\n"
                            f"I have successfully searched and synthesized data from **{len(self.sources)} websites**. "
                            f"The compiled, high-fidelity report file has been attached below.",
                    files=files_to_send
                )
                
                if os.path.exists(filename):
                    os.remove(filename)
                if os.path.exists("chart.png"):
                    os.remove("chart.png")
            except Exception as file_err:
                logger.error(f"Failed to write or transmit file payload: {file_err}")
                await bot._send_split_content(self.channel, report_markdown)
                
        self.status = "completed"
        await self.update_progress_msg(bot, active_idx=5, custom_header="✅ *Deep Research successfully completed! Report file delivered.*")


    async def search_duckduckgo(self, query: str) -> List[Dict[str, str]]:
        url = f"https://html.duckduckgo.com/html/?q={urllib.parse.quote(query)}"
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        }
        try:
            connector = aiohttp.TCPConnector(ssl=False)
            async with aiohttp.ClientSession(headers=headers, connector=connector) as session:
                async with session.get(url, timeout=12) as response:
                    if response.status != 200:
                        return []
                    html_data = await response.text()
                    
            soup = BeautifulSoup(html_data, "html.parser")
            results = []
            
            for node in soup.find_all("a", class_="result__snippet")[:12]:
                title_node = node.find_previous("a", class_="result__url")
                if title_node and title_node.get("href"):
                    title = title_node.get_text().strip()
                    link = title_node["href"]
                    snippet = node.get_text().strip()
                    
                    if link.startswith("//"):
                        link = "https:" + link
                    if "duckduckgo.com/l/" in link:
                        parsed = urllib.parse.urlparse(link)
                        params = urllib.parse.parse_qs(parsed.query)
                        if "uddg" in params:
                            link = params["uddg"][0]
                            
                    results.append({"url": link, "title": title, "snippet": snippet})
                    
            return results
        except Exception as e:
            logger.warning(f"DuckDuckGo query '{query}' failed: {e}")
            return []

    async def process_single_page(self, bot, url: str, title: str):
        if url in self.scraped_urls:
            return
        self.scraped_urls.add(url)
        
        logger.info(f"Deep Research scraping url: {url}")
        self.current_thought = f"Downloading and cleaning contents from: {url.split('//')[-1].split('/')[0]}..."
        try:
            markdown = await bot.link_reader.fetch_and_clean(url)
            if not markdown or markdown.startswith("[Error") or markdown.startswith("[Failed"):
                return
                
            self.sources.append({"url": url, "title": title})
            source_idx = len(self.sources)
            
            self.current_thought = f"Page downloaded. Running concurrent LLM fact extraction on: {title[:30]}..."
            
            async with self.extract_sem:
                await self.extract_facts_from_page(bot, title, url, markdown, source_idx)
                
            sub_links = re.findall(r'\[.*?\]\((https?://[^\s<>"]+?)\)', markdown)
            target_sub_links = []
            
            base_domain = url.split("//")[-1].split("/")[0]
            for link in sub_links:
                if len(target_sub_links) >= 2:
                    break
                if base_domain in link and link not in self.scraped_urls:
                    if any(kw in link.lower() for kw in ["pricing", "price", "hardware", "spec", "premium", "node"]):
                        target_sub_links.append(link)
                        
            if target_sub_links:
                logger.info(f"Subpage Crawler identified sub-links: {target_sub_links}")
                sub_crawl_tasks = [self.process_single_page(bot, sub_url, f"Nested Spec: {sub_url.split('/')[-1]}") for sub_url in target_sub_links]
                await asyncio.gather(*sub_crawl_tasks)
                
        except Exception as e:
            logger.warning(f"Error scraping page {url}: {e}")


    async def generate_queries(self, bot) -> List[str]:
        prompt = (
            f"You are the planning engine of a Deep Research Agent. Read the user's research request:\n"
            f"\"{self.primary_task}\"\n\n"
            f"Guidelines context:\n\"{self.guidelines or 'None'}\"\n\n"
            f"Generate a list of 4-6 distinct, highly targeted, and search-optimized keyword search queries "
            f"designed to uncover different facts, pricing tables, specifications, or community benchmarks of this topic.\n"
            f"Output ONLY a raw JSON string list. Do not write any explanations or markdown backticks.\n"
            f"Example format:\n[\"query 1\", \"query 2\", \"query 3\"]"
        )
        
        max_attempts = 3
        for attempt in range(max_attempts):
            try:
                response = await bot.chat_handler.client.aio.models.generate_content(
                    model=bot.chat_handler.premium_model,
                    contents=prompt,
                    config=types.GenerateContentConfig(temperature=0.2, response_mime_type="application/json")
                )
                text = response.text.strip()
                text_clean = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.IGNORECASE).strip()
                queries = json.loads(text_clean)
                if isinstance(queries, list) and queries:
                    return [q.strip() for q in queries if q.strip()]
            except Exception as err:
                logger.warning(f"LLM Planner query generation attempt {attempt + 1} failed: {err}")
                if attempt == max_attempts - 1:
                    break
                await asyncio.sleep(1.5 * (2 ** attempt))
            
        return [self.primary_task]

    async def filter_relevance(self, bot, candidates: List[Dict[str, str]], limit: int) -> List[Dict[str, str]]:
        if len(candidates) <= limit:
            return candidates
            
        repr_list = [{"idx": i, "title": c["title"], "url": c["url"], "snippet": c.get("snippet", "")} for i, c in enumerate(candidates)]
        
        prompt = (
            f"Read this list of candidate web pages found on the search indexes:\n"
            f"{json.dumps(repr_list, indent=2)}\n\n"
            f"Your task is to select the top {limit} most authoritative, factual, and relevant items "
            f"to scrape that will best answer the research topic: \"{self.primary_task}\".\n"
            f"Output ONLY a raw JSON list containing the integer indices of your selections.\n"
            f"Example format:\n[0, 2, 4, 9]"
        )
        
        max_attempts = 3
        for attempt in range(max_attempts):
            try:
                response = await bot.chat_handler.client.aio.models.generate_content(
                    model=bot.chat_handler.premium_model,
                    contents=prompt,
                    config=types.GenerateContentConfig(temperature=0.0, response_mime_type="application/json")
                )
                text = response.text.strip()
                text_clean = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.IGNORECASE).strip()
                indices = json.loads(text_clean)
                
                selected = []
                for idx in indices:
                    if 0 <= idx < len(candidates):
                        selected.append(candidates[idx])
                if selected:
                    return selected
            except Exception as e:
                logger.warning(f"Link relevance filtering attempt {attempt + 1} failed: {e}")
                if attempt == max_attempts - 1:
                    break
                await asyncio.sleep(1.5 * (2 ** attempt))
            
        return candidates[:limit]

    async def extract_facts_from_page(self, bot, title: str, url: str, markdown: str, source_idx: int):
        prompt = (
            f"You are the fact extraction worker of a Deep Research Agent.\n"
            f"We are researching: \"{self.primary_task}\"\n"
            f"Scraped Webpage: \"{title}\" (URL: {url})\n\n"
            f"Extract any specific, verified, non-trivial facts, pricing numbers, benchmarks, "
            f"hardware details, release dates, or community observations that help answer our topic.\n"
            f"Output ONLY a raw JSON list of strings containing these extracted facts. Keep them highly concise "
            f"and include numeric data where available. Do not include markdown formatting.\n"
            f"Example format:\n[\"Apex has a node in Virginia\", \"Host RAM starts at $4/GB\"]"
        )
        
        max_attempts = 3
        for attempt in range(max_attempts):
            try:
                response = await bot.chat_handler.client.aio.models.generate_content(
                    model=bot.chat_handler.premium_model,
                    contents=prompt,
                    config=types.GenerateContentConfig(temperature=0.2, response_mime_type="application/json")
                )
                text = response.text.strip()
                text_clean = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.IGNORECASE).strip()
                facts = json.loads(text_clean)
                
                if isinstance(facts, list):
                    for f in facts:
                        if f.strip():
                            self.fact_bank.append({"fact": f.strip(), "source_idx": source_idx})
                    break
            except Exception as e:
                logger.warning(f"Fact extraction attempt {attempt + 1} failed for {url}: {e}")
                if attempt == max_attempts - 1:
                    break
                await asyncio.sleep(1.5 * (2 ** attempt))

    async def analyze_gaps(self, bot) -> List[str]:
        facts_str = "\n".join([f"- [Source {f['source_idx']}] {f['fact']}" for f in self.fact_bank])
        
        prompt = (
            f"You are the self-correcting analyzer of a Deep Research Agent.\n"
            f"We are researching: \"{self.primary_task}\"\n\n"
            f"Here is our active bank of extracted facts so far:\n"
            f"{facts_str[:4000]}\n\n"
            f"Evaluate the database. Identify unresolved questions, critical details that are still unverified "
            f"(like specific East Coast latency figures, hidden server fees, hardware CPUs, or support benchmarks), "
            f"or discrepancies. Write 2-3 highly specific search queries targeting these missing details.\n"
            f"Output ONLY a raw JSON list of these search query strings. If you find no gaps and the facts are "
            f"thoroughly complete, output an empty list: []."
        )
        
        max_attempts = 3
        for attempt in range(max_attempts):
            try:
                response = await bot.chat_handler.client.aio.models.generate_content(
                    model=bot.chat_handler.premium_model,
                    contents=prompt,
                    config=types.GenerateContentConfig(temperature=0.0, response_mime_type="application/json")
                )
                text = response.text.strip()
                text_clean = re.sub(r'^```json\s*|\s*```$', '', text, flags=re.IGNORECASE).strip()
                queries = json.loads(text_clean)
                if isinstance(queries, list):
                    return [q.strip() for q in queries if q.strip()]
            except Exception as e:
                logger.warning(f"Gap analysis attempt {attempt + 1} failed: {e}")
                if attempt == max_attempts - 1:
                    break
                await asyncio.sleep(1.5 * (2 ** attempt))
            
        return []


    async def compile_report_sequential(self, bot) -> str:
        facts_str = "\n".join([f"- [Source {f['source_idx']}] {f['fact']}" for f in self.fact_bank])
        sources_str = "\n".join([f"- [Source {i+1}] Title: \"{s['title']}\" | URL: {s['url']}" for i, s in enumerate(self.sources)])
        
        system_instruction = (
            "You are an elite, professional research compiler and technical analyst. "
            "Write in a highly authoritative, objective, and dense reporting tone. "
            "Never use conversational intros, intros, or personal filler remarks."
        )
        
        ch1_prompt = (
            f"We are compiling an exhaustive research dossier on: \"{self.primary_task}\"\n"
            f"Flavor Focus: `{self.flavor.upper()}`\n\n"
            f"Write 'Chapter 1: Hardware Infrastructure & CPU Performance'. "
            f"Analyze and compare physical CPU architectures, SSD throughput levels, and core processor clocks. "
            f"Base your writeup strictly on this fact database, including inline footnotes linking back to the source indexes (e.g. [3]):\n"
            f"{facts_str[:4000]}\n\n"
            f"Write extensively, detailing every spec point found."
        )
        ch1_text = ""
        try:
            response = await bot.chat_handler.client.aio.models.generate_content(
                model=bot.chat_handler.premium_model, contents=ch1_prompt,
                config=types.GenerateContentConfig(system_instruction=system_instruction, temperature=0.4)
            )
            ch1_text = response.text.strip()
        except Exception as e:
            ch1_text = f"### Chapter 1: Hardware Infrastructure & CPU Performance\nFailed to draft chapter: {e}"

        ch2_prompt = (
            f"We are compiling an exhaustive research dossier on: \"{self.primary_task}\"\n"
            f"Here is our drafted Chapter 1:\n{ch1_text}\n\n"
            f"Write 'Chapter 2: US East Coast Network Latency & Connectivity'. "
            f"Detail pings, routing nodes, and regional performance. "
            f"Base your writeup strictly on this fact database, including inline footnotes (e.g. [1]):\n"
            f"{facts_str[:4000]}\n\n"
            f"Write extensively."
        )
        ch2_text = ""
        try:
            response = await bot.chat_handler.client.aio.models.generate_content(
                model=bot.chat_handler.premium_model, contents=ch2_prompt,
                config=types.GenerateContentConfig(system_instruction=system_instruction, temperature=0.4)
            )
            ch2_text = response.text.strip()
        except Exception as e:
            ch2_text = f"### Chapter 2: US East Coast Network Latency & Connectivity\nFailed to draft chapter: {e}"

        ch3_prompt = (
            f"We are compiling an exhaustive research dossier on: \"{self.primary_task}\"\n"
            f"Here is our drafted Chapter 2:\n{ch2_text}\n\n"
            f"Write 'Chapter 3: Optional Add-on Pricing, Dedicated IPs, and Target Use-Cases'. "
            f"Detail cost structures, waiving tiers, and explicit target user profiles. "
            f"Base your writeup strictly on this fact database, including inline footnotes (e.g. [8, 9]):\n"
            f"{facts_str[:4000]}\n\n"
            f"Write extensively."
        )
        ch3_text = ""
        try:
            response = await bot.chat_handler.client.aio.models.generate_content(
                model=bot.chat_handler.premium_model, contents=ch3_prompt,
                config=types.GenerateContentConfig(system_instruction=system_instruction, temperature=0.4)
            )
            ch3_text = response.text.strip()
        except Exception as e:
            ch3_text = f"### Chapter 3: Optional Add-on Pricing, Dedicated IPs, and Target Use-Cases\nFailed to draft chapter: {e}"

        complete_markdown = (
            f"## Analytical Research Dossier: {self.primary_task}\n"
            f"----------------------------------------\n\n"
            f"{ch1_text}\n\n"
            f"{ch2_text}\n\n"
            f"{ch3_text}\n\n"
            f"### 📚 Sources Bibliography\n"
            f"{sources_str}"
        )
        return complete_markdown