import discord
import logging
import asyncio
import json
import re
import os
import aiohttp
import urllib.parse
from bs4 import BeautifulSoup
from typing import Optional, List, Dict, Any, Tuple
from google.genai import types
from PIL import Image, ImageDraw
from datetime import datetime, timezone

import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from agents.base_agent import BaseAgentSession

logger = logging.getLogger("DeepResearchAgent")


def add_hyperlink(
    paragraph, url: str, text: str, color: str = "4183C4", underline: bool = True
):
    """
    Helper to inject a native, clickable hyperlink into a python-docx paragraph
    using low-level OpenXML schema configurations.
    """
    try:
        import docx.oxml as oxml
        import docx.oxml.ns as ns

        part = paragraph.part
        r_id = part.relate_to(
            url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )

        hyperlink = oxml.OxmlElement("w:hyperlink")
        hyperlink.set(oxml.ns.qn("r:id"), r_id)

        new_run = oxml.OxmlElement("w:r")
        rPr = oxml.OxmlElement("w:rPr")

        if color:
            c = oxml.OxmlElement("w:color")
            c.set(oxml.ns.qn("w:val"), color)
            rPr.append(c)

        if underline:
            u = oxml.OxmlElement("w:u")
            u.set(oxml.ns.qn("w:val"), "single")
            rPr.append(u)

        new_run.append(rPr)

        text_node = oxml.OxmlElement("w:t")
        text_node.text = text
        new_run.append(text_node)
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)
        return hyperlink
    except Exception as e:
        logger.warning(f"Could not inject XML hyperlink: {e}")
        paragraph.add_run(f" {text} ({url})")
        return None


def write_inline_formatted_text(paragraph, text: str):
    """
    Parses inline Markdown patterns (bold '**' and hyperlinks '[]()')
    and writes them natively as separate runs into the paragraph.
    """
    pattern = re.compile(r"(\*\*.*?\*\*|\[.*?\]\(.*?\))")
    parts = pattern.split(text)

    for part in parts:
        if not part:
            continue

        if part.startswith("**") and part.endswith("**"):
            clean_bold = part[2:-2]
            run = paragraph.add_run(clean_bold)
            run.bold = True

        elif part.startswith("[") and "](" in part and part.endswith(")"):
            link_match = re.match(r"\[(.*?)\]\((.*?)\)", part)
            if link_match:
                link_text, link_url = link_match.groups()
                add_hyperlink(paragraph, link_url, link_text)
            else:
                paragraph.add_run(part)

        else:
            paragraph.add_run(part)


def convert_markdown_to_docx(
    doc_instance, markdown_text: str, double_spaced: bool = False
):
    """
    Parses a markdown document line-by-line and appends native headings,
    bullet lists, and formatted paragraphs to a python-docx instance.
    Supports academic double spacing modifiers.
    """
    lines = markdown_text.split("\n")
    in_code_block = False

    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            continue

        if line_strip.startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc_instance.add_paragraph()
            if double_spaced:
                p.paragraph_format.line_spacing = 2.0
            run = p.add_run(line)
            run.font.name = "Consolas"
            continue

        if line_strip.startswith("# "):
            doc_instance.add_heading(line_strip[2:], level=1)
        elif line_strip.startswith("## "):
            doc_instance.add_heading(line_strip[3:], level=2)
        elif line_strip.startswith("### "):
            doc_instance.add_heading(line_strip[4:], level=3)

        elif line_strip.startswith("- ") or line_strip.startswith("* "):
            content = line_strip[2:]
            p = doc_instance.add_paragraph(style="List Bullet")
            if double_spaced:
                p.paragraph_format.line_spacing = 2.0
            write_inline_formatted_text(p, content)

        else:
            p = doc_instance.add_paragraph()
            if double_spaced:
                p.paragraph_format.line_spacing = 2.0
            write_inline_formatted_text(p, line_strip)


def generate_comparison_chart(
    fact_bank: List[Dict[str, Any]], filename: str = "chart.png"
) -> bool:
    """
    Dynamically scans the structured EAV fact bank, identifies comparative numeric
    attributes (such as cost, capacity, latency), and draws a beautifully proportioned
    dark-mode bar comparison chart without hardcoded brand matching.
    """
    logger.info("Initializing generalized visual comparison graph compile pass...")

    entity_metrics = {}

    price_pattern = re.compile(r"\$(\d+(?:\.\d+)?)")
    number_pattern = re.compile(r"\b(\d+)\s*(ms|gb|mb|tb|%)\b", re.IGNORECASE)

    detected_unit = ""

    for item in fact_bank:
        entity = item.get("entity", "").strip()
        fact = item.get("fact", "")
        if (
            not entity
            or not fact
            or entity.lower() == "none"
            or entity.lower() in ("various", "all")
        ):
            continue

        entity_title = entity.title()

        pr_match = price_pattern.search(fact)
        if pr_match:
            val = float(pr_match.group(1))
            if entity_title not in entity_metrics:
                entity_metrics[entity_title] = {"val": val, "unit": "$"}
                detected_unit = "$"
            continue

        num_match = number_pattern.search(fact)
        if num_match:
            val = float(num_match.group(1))
            unit = num_match.group(2).lower()
            if entity_title not in entity_metrics:
                entity_metrics[entity_title] = {"val": val, "unit": unit}
                detected_unit = unit
            continue

    if len(entity_metrics) < 2:
        logger.info(
            "Skipping chart generation: Insufficient comparative datasets discovered."
        )
        return False

    providers = list(entity_metrics.keys())[:5]

    width, height = 800, 500
    img = Image.new("RGB", (width, height), "#1e1f22")
    draw = ImageDraw.Draw(img)

    unit_label = (
        f"({detected_unit.upper()})" if detected_unit and detected_unit != "$" else ""
    )
    draw.text(
        (30, 20),
        f"Web Research Comparative Data Matrix {unit_label}".strip(),
        fill="#ffffff",
    )
    draw.line([(30, 45), (770, 45)], fill="#2f3136", width=2)

    y_start = 80
    row_height = 80

    max_val = max(item["val"] for item in entity_metrics.values())
    if max_val <= 0:
        max_val = 1.0

    for i, prov in enumerate(providers):
        y_pos = y_start + (i * row_height)
        data = entity_metrics[prov]
        val = data["val"]
        unit = data["unit"]

        draw.text((30, y_pos), prov[:18], fill="#ffffff")

        bar_width = min(400, int((val / max_val) * 400))
        bar_width = max(10, bar_width)

        color_hex = "#248046" if unit == "$" else "#5865f2"

        draw.rectangle([(180, y_pos), (180 + bar_width, y_pos + 22)], fill=color_hex)

        display_str = f"${val:.2f}" if unit == "$" else f"{int(val)}{unit}"
        draw.text((190 + bar_width, y_pos + 4), display_str, fill="#949ba4")

    try:
        img.save(filename)
        return True
    except Exception as img_err:
        logger.error(f"Failed to compile visual Pillow chart: {img_err}")
        return False


class DeepResearchSession(BaseAgentSession):
    def __init__(
        self,
        thread_id: int,
        user_id: int,
        prompt: str,
        loaded_contexts: str,
        channel: discord.Thread,
        depth: str = "standard",
        format_type: str = "markdown",
        flavor: str = "executive",
        domains_filter: str = "",
        guidelines: str = "",
    ):
        super().__init__(thread_id, user_id, prompt, loaded_contexts, channel)
        self.depth = depth.lower().strip()
        self.format_type = format_type.lower().strip()
        self.flavor = flavor.lower().strip()
        self.domains_filter = domains_filter.strip()
        self.guidelines = guidelines.strip()

        self.sources: List[Dict[str, str]] = []
        self.fact_bank: List[Dict[str, Any]] = []
        self.scraped_urls: set = set()

        self.clean_title: str = "Analytical Research Dossier"

        self.chapters_outline: Dict[str, Dict[str, str]] = {}

        self.partial_drafts: Dict[str, str] = {}

        self.works_cited_list: List[str] = []

        self.extract_sem = asyncio.Semaphore(3)

        self.plan_text: str = "Planning research targets..."
        self.current_thought: str = "Initiating active plan generation pass..."
        self.progress_msg: Optional[discord.Message] = None
        self.checklist_lines: List[str] = [
            "⚪ **Step 1: Planning dynamic report outline...**",
            "⚪ **Step 2: Crawling web indexes...**",
            "⚪ **Step 3: Scraping relevant pages & extracting EAV facts...**",
            "⚪ **Step 4: Performing recursive discrepancy reconciliation...**",
            "⚪ **Step 5: Compiling and formatting analytical report...**",
        ]

    def compile_react_transcript(self) -> str:
        return f"Deep Research Active. Depth: {self.depth}. Targets scraped: {len(self.scraped_urls)}."

    async def update_progress_msg(self, bot, active_idx: int, custom_header: str = ""):
        """Edits the active status checklist message in the thread using Components V2."""
        from agents.router_views import ViewResearchButton

        for i in range(len(self.checklist_lines)):
            line_clean = re.sub(r"^[⚪⏳✅]\s*", "", self.checklist_lines[i])
            if i < active_idx:
                self.checklist_lines[i] = f"✅ {line_clean}"
            elif i == active_idx:
                self.checklist_lines[i] = f"⏳ {line_clean}"
            else:
                self.checklist_lines[i] = f"⚪ {line_clean}"

        title_icon = "🔎" if active_idx < 4 else "📋"
        title_label = (
            "Deep Research in Progress" if active_idx < 4 else "Research Concluding"
        )

        card_content = (
            f"### {title_icon} {title_label}\n"
            f"----------------------------------------\n"
            f"{custom_header}\n\n"
            f"**Execution Checklist:**\n" + "\n".join(self.checklist_lines)
        )

        view = discord.ui.View(timeout=None)
        view.add_item(ViewResearchButton(self))

        if not self.progress_msg:
            self.progress_msg = await self.channel.send(content=card_content, view=view)
        else:
            try:
                await self.progress_msg.edit(content=card_content, view=view)
            except Exception as edit_err:
                logger.warning(f"Could not edit progress message: {edit_err}")

    async def execute_tick(self, bot):
        """Asynchronous execution controller driving the Deep Research engine."""
        try:
            await self.run_research(bot)
        except Exception as e:
            logger.error(f"Deep Research Loop crashed: {e}", exc_info=True)
            self.status = "paused_error"
            await self.channel.send(
                f"🛑 **Research Interrupted:** An unexpected crash occurred during synthesis mapping: `{e}`"
            )

    async def run_research(self, bot):
        """Master orchestrator conducting Planning, Parallel search/scrape, Gaps checks, and Synthesis."""
        logger.info(
            f"Starting Deep Research loop in thread {self.thread_id}. Task: '{self.primary_task[:30]}'"
        )

        await self.update_progress_msg(
            bot,
            active_idx=0,
            custom_header="*Formulating deconstruction plan and targeting sub-questions...*",
        )
        queries = await self.generate_queries(bot)

        await self.generate_custom_outline(bot)

        self.plan_text = (
            f"Drafted {len(queries)} initial query targets across public indexes:\n"
            + "\n".join([f"- {q}" for q in queries])
        )
        self.current_thought = "Executing parallel search queries..."

        await self.update_progress_msg(
            bot,
            active_idx=1,
            custom_header=f"📡 *Crawling web indexes concurrently using `{len(queries)}` target queries...*",
        )

        search_tasks = [self.search_duckduckgo(query) for query in queries]
        search_results = await asyncio.gather(*search_tasks)

        candidates = []
        seen_links = set()

        exclude_domains = []
        if self.domains_filter:
            exclude_domains = [
                d.replace("-", "").strip().lower()
                for d in self.domains_filter.split(",")
                if d.strip().startswith("-")
            ]

        for res_list in search_results:
            for item in res_list:
                url = item.get("url", "")
                if url in seen_links or url in self.scraped_urls:
                    continue

                domain_lower = url.lower()
                if any(ex in domain_lower for ex in exclude_domains if ex):
                    continue

                seen_links.add(url)
                candidates.append(item)

        if not candidates:
            candidates = await self.search_duckduckgo(self.primary_task)

        await self.update_progress_msg(
            bot,
            active_idx=2,
            custom_header=f"📄 *Relevance filtering candidates. Selecting top sources...*",
        )

        max_links = (
            5
            if self.depth == "brief"
            else (
                15
                if self.depth == "standard"
                else 30 if self.depth == "exhaustive" else 60
            )
        )

        candidate_slice_limit = (
            15
            if self.depth == "brief"
            else (
                45
                if self.depth == "standard"
                else 75 if self.depth == "exhaustive" else 100
            )
        )

        selected_candidates = await self.filter_relevance(
            bot, candidates[:candidate_slice_limit], limit=max_links
        )

        self.current_thought = f"Selected {len(selected_candidates)} highly relevant pages to scrape. Extracting EAV structured parameters..."
        await self.update_progress_msg(
            bot,
            active_idx=2,
            custom_header=f"📄 *Concurrently scraping `{len(selected_candidates)}` pages and analyzing contents...*",
        )

        scrape_tasks = [
            self.process_single_page(bot, cand.get("url"), cand.get("title", ""))
            for cand in selected_candidates
        ]
        await asyncio.gather(*scrape_tasks)

        self.current_thought = (
            "Checking database for conflicting measurements and data discrepancies..."
        )
        await self.update_progress_msg(
            bot,
            active_idx=3,
            custom_header="🔍 *Identifying discrepancies and resolving conflicts inside fact databases...*",
        )

        await self.reconcile_contradictions(bot)

        if self.depth != "brief" and len(self.sources) > 0:
            gap_queries = await self.analyze_gaps(bot)
            if gap_queries:
                self.current_thought = f"Discovered knowledge gaps. Executing secondary research loop: {gap_queries}"
                await self.update_progress_msg(
                    bot,
                    active_idx=3,
                    custom_header=f"📡 *Recursive pass: Scraping `{len(gap_queries)}` secondary targets...*",
                )

                gap_search_tasks = [self.search_duckduckgo(g_q) for g_q in gap_queries]
                gap_search_results = await asyncio.gather(*gap_search_tasks)

                gap_candidates = []
                for res_list in gap_search_results:
                    for item in res_list:
                        url = item.get("url", "")
                        if url not in self.scraped_urls and url not in [
                            c.get("url") for c in gap_candidates
                        ]:
                            gap_candidates.append(item)

                if gap_candidates:
                    gap_scrape_tasks = [
                        self.process_single_page(
                            bot, cand.get("url"), cand.get("title", "")
                        )
                        for cand in gap_candidates[:4]
                    ]
                    await asyncio.gather(*gap_scrape_tasks)
                    await self.reconcile_contradictions(bot)

        if self.format_type == "mla":
            self.current_thought = (
                "Executing isolated scholarly bibliography formatting pass..."
            )
            await self.update_progress_msg(
                bot,
                active_idx=3,
                custom_header="📚 *Compiling structured MLA Works Cited index...*",
            )
            await self.compile_works_cited_mla(bot)

        self.current_thought = "Designing custom chapter outline and drafting report chapters sequentially..."
        await self.update_progress_msg(
            bot,
            active_idx=4,
            custom_header="✍️ *Synthesizing collected fact indexes and generating formatted document...*",
        )

        report_markdown = await self.compile_report_sequential(bot)

        chart_generated = generate_comparison_chart(self.fact_bank, "chart.png")

        slug = re.sub(r"[^a-zA-Z0-9]", "_", self.clean_title).strip("_")

        if self.format_type in ("docx", "mla"):
            filename = f"Research_Report_{slug}.docx"
            try:
                import docx
                from docx.shared import Inches, Pt

                doc = docx.Document()

                if self.format_type == "mla":

                    style = doc.styles["Normal"]
                    font = style.font
                    font.name = "Times New Roman"
                    font.size = Pt(12)

                    for section in doc.sections:
                        section.top_margin = Inches(1)
                        section.bottom_margin = Inches(1)
                        section.left_margin = Inches(1)
                        section.right_margin = Inches(1)

                    p_name = doc.add_paragraph()
                    p_name.paragraph_format.line_spacing = 2.0
                    p_name.add_run("Research Analyst")

                    p_inst = doc.add_paragraph()
                    p_inst.paragraph_format.line_spacing = 2.0
                    p_inst.add_run("Deep Research Intelligence Assistant")

                    p_course = doc.add_paragraph()
                    p_course.paragraph_format.line_spacing = 2.0
                    p_course.add_run("Global Strategic Intelligence")

                    p_date = doc.add_paragraph()
                    p_date.paragraph_format.line_spacing = 2.0
                    current_date_str = datetime.now(timezone.utc).strftime("%d %B %Y")
                    p_date.add_run(current_date_str)

                    p_title = doc.add_paragraph()
                    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_title.paragraph_format.line_spacing = 2.0
                    run_t = p_title.add_run(self.clean_title)
                    run_t.bold = True
                else:

                    doc.add_heading(f"Research Dossier: {self.clean_title}", level=0)

                if chart_generated and self.format_type != "mla":
                    doc.add_heading("Visual Metric Graph Matrix", level=1)
                    doc.add_picture("chart.png", width=docx.shared.Inches(5.5))
                    doc.add_paragraph(
                        "Figure 1.0: Compiled specifications visual matrix."
                    )

                convert_markdown_to_docx(
                    doc, report_markdown, double_spaced=(self.format_type == "mla")
                )

                if self.format_type == "mla" and self.works_cited_list:
                    doc.add_page_break()
                    p_wc_title = doc.add_paragraph()
                    p_wc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_wc_title.paragraph_format.line_spacing = 2.0
                    run_wc = p_wc_title.add_run("Works Cited")
                    run_wc.bold = True

                    for citation in self.works_cited_list:
                        p_cit = doc.add_paragraph()
                        p_cit.paragraph_format.line_spacing = 2.0
                        write_inline_formatted_text(p_cit, citation)

                doc.save(filename)
                file = discord.File(fp=filename, filename=filename)

                files_to_send = [file]
                if chart_generated and self.format_type != "mla":
                    files_to_send.append(
                        discord.File(fp="chart.png", filename="chart.png")
                    )

                mention_prefix = (
                    f"🛎️ <@{self.user_id}>, \n" if self.depth == "extreme" else ""
                )

                await self.channel.send(
                    content=f"{mention_prefix}### 📋 Deep Research Summary Report Completed!\n"
                    f"I have successfully searched and compiled data from **{len(self.sources)} websites** "
                    f"into a polished Word document (.docx). The file is attached below.",
                    files=files_to_send,
                )

                if os.path.exists(filename):
                    os.remove(filename)
                if os.path.exists("chart.png"):
                    os.remove("chart.png")

            except Exception as docx_err:
                logger.error(
                    f"Failed compiling Word Document: {docx_err}. Falling back to standard Markdown compilation..."
                )
                self.format_type = "markdown"

        if self.format_type == "markdown":
            filename = f"Research_Report_{slug}.md"
            try:
                with open(filename, "w", encoding="utf-8") as f:
                    f.write(report_markdown)

                file = discord.File(fp=filename, filename=filename)
                files_to_send = [file]
                if chart_generated:
                    files_to_send.append(
                        discord.File(fp="chart.png", filename="chart.png")
                    )

                mention_prefix = (
                    f"🛎️ <@{self.user_id}>, \n" if self.depth == "extreme" else ""
                )

                await self.channel.send(
                    content=f"{mention_prefix}### 📋 Deep Research Summary Report Completed!\n"
                    f"I have successfully searched and synthesized data from **{len(self.sources)} websites**. "
                    f"The compiled, high-fidelity report file has been attached below.",
                    files=files_to_send,
                )

                if os.path.exists(filename):
                    os.remove(filename)
                if os.path.exists("chart.png"):
                    os.remove("chart.png")
            except Exception as file_err:
                logger.error(f"Failed to write or transmit file payload: {file_err}")
                await bot._send_split_content(self.channel, report_markdown)

        self.status = "completed"
        await self.update_progress_msg(
            bot,
            active_idx=5,
            custom_header="✅ *Deep Research successfully completed! Report file delivered.*",
        )

    async def search_duckduckgo(self, query: str) -> List[Dict[str, str]]:
        """High-performance async DuckDuckGo crawler returning candidates (title, url, snippet)."""
        url = f"https://html.duckduckgo.com/html/?q={urllib.parse.quote(query)}"
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        }
        try:
            connector = aiohttp.TCPConnector(ssl=False)
            async with aiohttp.ClientSession(
                headers=headers, connector=connector
            ) as session:
                async with session.get(url, timeout=12) as response:
                    if response.status != 200:
                        return []
                    html_data = await response.text()

            soup = BeautifulSoup(html_data, "html.parser")
            results = []

            for node in soup.find_all("a", class_="result__snippet")[:12]:
                title_node = node.find_previous("a", class_="result__url")
                if title_node and title_node.get("href"):
                    title = title_node.get_text().strip()
                    link = title_node["href"]
                    snippet = node.get_text().strip()

                    if link.startswith("//"):
                        link = "https:" + link
                    if "duckduckgo.com/l/" in link:
                        parsed = urllib.parse.urlparse(link)
                        params = urllib.parse.parse_qs(parsed.query)
                        if "uddg" in params:
                            link = params["uddg"][0]

                    results.append({"url": link, "title": title, "snippet": snippet})

            return results
        except Exception as e:
            logger.warning(f"DuckDuckGo query '{query}' failed: {e}")
            return []

    async def process_single_page(self, bot, url: str, title: str):
        """Scrapes a page and extracts its facts under a strict concurrency throttle."""
        if url in self.scraped_urls:
            return
        self.scraped_urls.add(url)

        logger.info(f"Deep Research scraping url: {url}")
        self.current_thought = f"Downloading and cleaning contents from: {url.split('//')[-1].split('/')[0]}..."
        try:
            markdown = await bot.link_reader.fetch_and_clean(url)
            if (
                not markdown
                or markdown.startswith("[Error")
                or markdown.startswith("[Failed")
            ):
                return

            self.sources.append({"url": url, "title": title})
            source_idx = len(self.sources)

            self.current_thought = f"Page downloaded. Running concurrent LLM fact extraction on: {title[:30]}..."

            async with self.extract_sem:
                await self.extract_facts_from_page(
                    bot, title, url, markdown, source_idx
                )

            sub_links = re.findall(r'\[.*?\]\((https?://[^\s<>"]+?)\)', markdown)
            target_sub_links = []

            base_domain = url.split("//")[-1].split("/")[0]
            for link in sub_links:
                if len(target_sub_links) >= 2:
                    break
                if base_domain in link and link not in self.scraped_urls:
                    if any(
                        kw in link.lower()
                        for kw in [
                            "pricing",
                            "price",
                            "hardware",
                            "spec",
                            "premium",
                            "node",
                        ]
                    ):
                        target_sub_links.append(link)

            if target_sub_links:
                sub_crawl_tasks = [
                    self.process_single_page(
                        bot, sub_url, f"Nested Spec: {sub_url.split('/')[-1]}"
                    )
                    for sub_url in target_sub_links
                ]
                await asyncio.gather(*sub_crawl_tasks)

        except Exception as e:
            logger.warning(f"Error scraping page {url}: {e}")

    async def generate_queries(self, bot) -> List[str]:
        """Planning phase: Deconstruct prompt and generate search terms."""
        prompt = (
            f"You are the planning engine of a Deep Research Agent. Read the user's research request:\n"
            f'"{self.primary_task}"\n\n'
            f"Guidelines context:\n\"{self.guidelines or 'None'}\"\n\n"
            f"Generate a list of 4-6 distinct, highly targeted, and search-optimized keyword search queries "
            f"designed to uncover different facts, specifications, or comparisons on this topic.\n"
            f"Output ONLY a raw JSON string list. Do not write any explanations or markdown backticks.\n"
            f'Example format:\n["query 1", "query 2", "query 3"]'
        )

        max_attempts = 3
        for attempt in range(max_attempts):
            try:
                response = await bot.chat_handler.client.aio.models.generate_content(
                    model=bot.chat_handler.premium_model,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        temperature=0.2, response_mime_type="application/json"
                    ),
                )
                text = (response.text or "").strip()
                text_clean = re.sub(
                    r"^```json\s*|\s*```$", "", text, flags=re.IGNORECASE
                ).strip()
                queries = json.loads(text_clean)
                if isinstance(queries, list) and queries:
                    return [q.strip() for q in queries if q.strip()]
            except Exception as err:
                logger.warning(
                    f"LLM Planner query generation attempt {attempt + 1} failed: {err}"
                )
                if attempt == max_attempts - 1:
                    break
                await asyncio.sleep(1.5 * (2**attempt))

        return [self.primary_task]

    async def generate_custom_outline(self, bot):
        """
        Planning Phase: Creates an optimized dynamic chapter outline customized specifically
        to the user's query, scaling the total chapter count directly based on research Depth.
        Pre-populates placeholder files so they are immediately readable inside the 'Live Draft' tab.
        """

        max_chapters = (
            2
            if self.depth == "brief"
            else (
                3
                if self.depth == "standard"
                else 4 if self.depth == "exhaustive" else 5
            )
        )

        layout_flavor_guide = (
            "academic paper essay structure"
            if self.format_type == "mla"
            else f"corporate `{self.flavor.upper()}` structure"
        )

        outline_prompt = (
            f'You are the lead editor of a Deep Research Agent. We are researching: "{self.primary_task}"\n'
            f"Formatting Flavor requested: `{self.flavor.upper()}`\n\n"
            f"Create a structural, customized {max_chapters}-chapter report outline conforming to a {layout_flavor_guide}. "
            f"Each chapter should address a separate structural facet of the topic.\n"
            f"You must also extract a clean, concise, 2-4 word Title for this dossier (e.g., 'Quantum Security Threat') "
            f"that captures the core theme of the prompt without including instructions.\n"
            f"Output a JSON object with this exact schema:\n"
            f"{{\n"
            f'  "title": "Consolidated Dossier Title",\n'
            f'  "outline": {{\n'
            f'     "Chapter 1": {{\n'
            f'        "title": "Title of Chapter 1",\n'
            f'        "objective": "Analytical objective"\n'
            f"     }},\n"
            f"     ...\n"
            f'     "Chapter {max_chapters}": {{\n'
            f'        "title": "Title of Chapter {max_chapters}",\n'
            f'        "objective": "Analytical objective"\n'
            f"     }}\n"
            f"  }}\n"
            f"}}\n\n"
            f"Output ONLY raw valid JSON matching the schema, with no markdown wrapping blocks."
        )

        default_outline = {}
        for idx in range(1, max_chapters + 1):
            default_outline[f"Chapter {idx}"] = {
                "title": f"Chapter {idx}: Research Analysis Segment",
                "objective": "Evaluate key metrics and structured datasets related to the query.",
            }

        self.chapters_outline = default_outline
        self.clean_title = "Analytical Research Dossier"

        try:
            response = await bot.chat_handler.client.aio.models.generate_content(
                model=bot.chat_handler.premium_model,
                contents=outline_prompt,
                config=types.GenerateContentConfig(
                    temperature=0.2, response_mime_type="application/json"
                ),
            )
            out_clean = re.sub(
                r"^```json\s*|\s*```$",
                "",
                (response.text or "").strip(),
                flags=re.IGNORECASE,
            ).strip()
            parsed = json.loads(out_clean)
            if isinstance(parsed, dict) and len(parsed) > 0:
                self.clean_title = parsed.get(
                    "title", "Analytical Research Dossier"
                ).strip()
                self.chapters_outline = parsed.get("outline", self.chapters_outline)
        except Exception as outline_err:
            logger.warning(
                f"Dynamic outline generation failed (applying fallback placeholders): {outline_err}"
            )

        self.partial_drafts.clear()
        for ch_key, ch_spec in self.chapters_outline.items():
            title = ch_spec.get("title", ch_key)
            obj = ch_spec.get("objective", "Evaluate contextual specs.")
            self.partial_drafts[title] = (
                f"*Status: Outline planned and locked. Waiting for web harvesting and fact extraction to complete...*\n\n"
                f"**Planned Objective:** {obj}"
            )

    async def filter_relevance(
        self, bot, candidates: List[Dict[str, str]], limit: int
    ) -> List[Dict[str, str]]:
        """Lightweight LLM call filtering search snippets to select the top relevant pages."""
        if len(candidates) <= limit:
            return candidates

        repr_list = [
            {
                "idx": i,
                "title": c["title"],
                "url": c["url"],
                "snippet": c.get("snippet", ""),
            }
            for i, c in enumerate(candidates)
        ]

        prompt = (
            f"Read this list of candidate web pages found on the search indexes:\n"
            f"{json.dumps(repr_list, indent=2)}\n\n"
            f"Your task is to select the top {limit} most authoritative, factual, and relevant items "
            f'to scrape that will best answer the research topic: "{self.primary_task}".\n'
            f"Output ONLY a raw JSON list containing the integer indices of your selections.\n"
            f"Example format:\n[0, 2, 4, 9]"
        )

        max_attempts = 3
        for attempt in range(max_attempts):
            try:
                response = await bot.chat_handler.client.aio.models.generate_content(
                    model=bot.chat_handler.premium_model,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        temperature=0.0, response_mime_type="application/json"
                    ),
                )
                text = (response.text or "").strip()
                text_clean = re.sub(
                    r"^```json\s*|\s*```$", "", text, flags=re.IGNORECASE
                ).strip()
                indices = json.loads(text_clean)

                selected = []
                for idx in indices:
                    if 0 <= idx < len(candidates):
                        selected.append(candidates[idx])
                if selected:
                    return selected
            except Exception as e:
                logger.warning(
                    f"Link relevance filtering attempt {attempt + 1} failed: {e}"
                )
                if attempt == max_attempts - 1:
                    break
                await asyncio.sleep(1.5 * (2**attempt))

        return candidates[:limit]

    async def extract_facts_from_page(
        self, bot, title: str, url: str, markdown: str, source_idx: int
    ):
        """Concurrent LLM call: extracts structured, semantic Entity-Attribute-Value facts with citation metadata."""
        prompt = (
            f"You are the structured fact extraction worker of a Deep Research Agent.\n"
            f'We are researching: "{self.primary_task}"\n'
            f'Scraped Webpage: "{title}" (URL: {url})\n\n'
            f"Extract specific, verified facts, metrics, pricing scales, attributes, or performance benchmarks.\n"
            f"You MUST output raw JSON following this schema list:\n"
            f"[\n"
            f"  {{\n"
            f'    "entity": "The target brand, competitor, or subject name (e.g. PebbleHost, Toyota)",\n'
            f'    "attribute": "The specific feature, capacity, or plan parameter evaluated (e.g. entry-price, battery-size)",\n'
            f'    "value": "The precise scalar, text, or dollar metric discovered",\n'
            f'    "fact": "A complete factual contextual statement describing this parameter in detail."\n'
            f"  }}\n"
            f"]\n\n"
            f"Output ONLY raw valid JSON, no explanations, no wrapping except the JSON block."
        )

        max_attempts = 3
        for attempt in range(max_attempts):
            try:
                response = await bot.chat_handler.client.aio.models.generate_content(
                    model=bot.chat_handler.premium_model,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        temperature=0.1, response_mime_type="application/json"
                    ),
                )
                text = (response.text or "").strip()
                text_clean = re.sub(
                    r"^```json\s*|\s*```$", "", text, flags=re.IGNORECASE
                ).strip()
                facts = json.loads(text_clean)

                if isinstance(facts, list):
                    for f in facts:
                        if isinstance(f, dict) and f.get("fact"):
                            self.fact_bank.append(
                                {
                                    "entity": f.get("entity", "None").strip(),
                                    "attribute": f.get("attribute", "None").strip(),
                                    "value": f.get("value", "None").strip(),
                                    "fact": f["fact"].strip(),
                                    "source_idx": source_idx,
                                }
                            )
                    break
            except Exception as e:
                logger.warning(
                    f"Fact extraction attempt {attempt + 1} failed for {url}: {e}"
                )
                if attempt == max_attempts - 1:
                    break
                await asyncio.sleep(1.5 * (2**attempt))

    async def reconcile_contradictions(self, bot):
        """Pre-Synthesis check: maps existing fact bank, identifies conflicting statements, and resolves them."""
        if len(self.fact_bank) < 2:
            return

        facts_summary = []
        for idx, item in enumerate(self.fact_bank):
            facts_summary.append(
                {
                    "index": idx,
                    "entity": item.get("entity"),
                    "attribute": item.get("attribute"),
                    "value": item.get("value"),
                    "fact": item.get("fact"),
                }
            )

        prompt = (
            f"You are the contradiction resolution engine of a Deep Research Agent.\n"
            f"Analyze this database of extracted research parameters:\n"
            f"{json.dumps(facts_summary[:60], indent=2)}\n\n"
            f"Identify any conflicting entries (e.g. two separate values, specs, or pricing rules claimed for the "
            f"same entity attribute). Determine which claims should be removed or marked as stale/outdated "
            f"(such as prioritizing 2026 data over 2024 information or official sources over hearsay).\n\n"
            f"Output a JSON list of integer indices that should be purged from our database because they are "
            f"invalid, contradictory, or obsolete. Example:\n"
            f"[4, 12, 18]\n\n"
            f"If there are no contradictions, output: []"
        )
        try:
            response = await bot.chat_handler.client.aio.models.generate_content(
                model=bot.chat_handler.premium_model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.0, response_mime_type="application/json"
                ),
            )
            text = (response.text or "").strip()
            text_clean = re.sub(
                r"^```json\s*|\s*```$", "", text, flags=re.IGNORECASE
            ).strip()
            purge_indices = json.loads(text_clean)

            if isinstance(purge_indices, list) and purge_indices:
                logger.info(
                    f"Reconciliation engine identified contradictions. Removing indices: {purge_indices}"
                )
                self.fact_bank = [
                    item
                    for i, item in enumerate(self.fact_bank)
                    if i not in purge_indices
                ]
        except Exception as e:
            logger.warning(
                f"Contradiction check failed (skipping resolution pass): {e}"
            )

    async def compile_works_cited_mla(self, bot):
        """
        Dedicated Step 4.5: Academic MLA Bibliography compiler.
        Translates raw scraped webpage sources into alphabetically-sorted, MLA 9th Edition citations.
        """
        source_records = []
        for idx, s in enumerate(self.sources):

            domain = s["url"].split("//")[-1].split("/")[0].replace("www.", "")
            source_records.append(
                {
                    "source_idx": idx + 1,
                    "title": s["title"],
                    "url": s["url"],
                    "domain_publisher": domain,
                    "access_date": datetime.now(timezone.utc).strftime("%d %b. %Y"),
                }
            )

        prompt = (
            f"You are an academic research librarian specializing in MLA 9th Edition formatting.\n"
            f"Compile a formal, academically rigorous Works Cited bibliography for the following sources:\n"
            f"{json.dumps(source_records, indent=2)}\n\n"
            f"Rules:\n"
            f"1. Generate a valid MLA 9th edition citation string for EVERY source. Guess authors if a prominent "
            f"publication organization represents the writer, or use domain titles.\n"
            f"2. You must output the links using standard markdown linking formatting `[Domain](URL)` inside the citation string.\n"
            f"3. Sort the citations alphabetically by author's last name or publisher title.\n"
            f"4. Output ONLY a raw JSON string list of these formatted citations. Do not write markdown wrapping blocks.\n"
            f"Example:\n"
            f"[\n"
            f'  "Ivezic, Marin. \\"Q-Day Revisited.\\" *LinkedIn*, May 2026, [LinkedIn](https://linkedin.com/...)",\n'
            f'  "Palo Alto Networks. \\"PQC Standards Overview.\\" *Palo Alto*, Jan. 2026, [Palo Alto](https://palo...)"\n'
            f"]"
        )
        try:
            response = await bot.chat_handler.client.aio.models.generate_content(
                model=bot.chat_handler.premium_model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.1, response_mime_type="application/json"
                ),
            )
            out_clean = re.sub(
                r"^```json\s*|\s*```$",
                "",
                (response.text or "").strip(),
                flags=re.IGNORECASE,
            ).strip()
            citations_list = json.loads(out_clean)
            if isinstance(citations_list, list):
                self.works_cited_list = citations_list
                logger.info(
                    f"Scholarly citations compiled successfully. Works Cited count: {len(self.works_cited_list)}"
                )
        except Exception as e:
            logger.error(f"Scholarly MLA works cited compilation pass failed: {e}")

            self.works_cited_list = [
                f"\"Title: {s['title']}.\" Web URL: [Link]({s['url']})"
                for s in self.sources
            ]

    async def analyze_gaps(self, bot) -> List[str]:
        """Gap analysis check to write targeted search lookups."""
        facts_str = "\n".join(
            [f"- [Source {f['source_idx']}] {f['fact']}" for f in self.fact_bank]
        )

        prompt = (
            f"You are the self-correcting analyzer of a Deep Research Agent.\n"
            f'We are researching: "{self.primary_task}"\n\n'
            f"Here is our active database of extracted facts:\n"
            f"{facts_str[:4000]}\n\n"
            f"Determine if any crucial objectives or target details remain unverified. "
            f"Write 2-3 specific search queries to resolve missing details.\n"
            f"Output ONLY a raw JSON list of queries, or [] if complete."
        )

        max_attempts = 3
        for attempt in range(max_attempts):
            try:
                response = await bot.chat_handler.client.aio.models.generate_content(
                    model=bot.chat_handler.premium_model,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        temperature=0.0, response_mime_type="application/json"
                    ),
                )
                text = (response.text or "").strip()
                text_clean = re.sub(
                    r"^```json\s*|\s*```$", "", text, flags=re.IGNORECASE
                ).strip()
                queries = json.loads(text_clean)
                if isinstance(queries, list):
                    return [q.strip() for q in queries if q.strip()]
            except Exception as e:
                logger.warning(f"Gap check attempt {attempt + 1} failed: {e}")
                if attempt == max_attempts - 1:
                    break
                await asyncio.sleep(1.5 * (2**attempt))

        return []

    async def compile_report_sequential(self, bot) -> str:
        """Runs the custom progressive drafting pipeline to build deep-dive dossiers."""
        system_instruction = (
            "You are an elite research compiler and technical academic. "
            "Write in a highly authoritative, objective, and dense reporting tone. "
            "Never use conversational intros, greetings, or personal filler remarks. "
        )

        if self.format_type == "mla":
            system_instruction += (
                "You are writing a scholarly research paper in strict MLA format. "
                "You MUST use double spacing, write with professional vocabulary, and use parenthetical inline "
                "citations matching the source's publisher or domain name (e.g. write '(Sustainableatlas)' or "
                "'(Carbonfact)' instead of numbered indexes like '[3]' or '[Source 3]'). "
                "The source facts you are provided with are prefix-wrapped with their dynamic citation keys, "
                "such as `(Sustainableatlas) | ...` or `(Nist) | ...`. You MUST extract and use these exact "
                "parenthetical inline keys in your text (e.g. write '...emissions (Sustainableatlas).' or "
                "'...standards (Nist).') when citing claims. Do not write a 'Sources' or 'References' section "
                "at the end—the Works Cited bibliography is compiled separately."
            )
        else:
            system_instruction += (
                "Always include numbered citation footnotes pointing back to source indexes (e.g. [3]) "
                "matching the fact database."
            )

        def filter_relevant_facts(chapter_goal: str) -> str:
            facts_list = []

            def get_clean_publisher(source_idx: int) -> str:
                if 0 < source_idx <= len(self.sources):
                    url = self.sources[source_idx - 1]["url"]
                    domain = url.split("//")[-1].split("/")[0].replace("www.", "")
                    parts = domain.split(".")

                    name = (
                        parts[0]
                        if parts[0] not in ("blog", "docs", "news", "journals", "pmc")
                        else parts[1]
                    )
                    return name.title()
                return f"Source {source_idx}"

            for f in self.fact_bank:
                ent = f.get("entity", "")
                attr = f.get("attribute", "")
                fact = f.get("fact", "")
                src = f.get("source_idx", 0)

                keywords = re.findall(
                    r"\b\w{4,}\b",
                    chapter_goal.lower() + " " + ent.lower() + " " + attr.lower(),
                )
                if any(
                    kw in fact.lower()
                    for kw in keywords
                    if kw not in ("about", "with", "this", "from", "that")
                ):
                    pub = get_clean_publisher(src)
                    facts_list.append(f"- ({pub}) | {ent} - {attr}: {fact}")

            if not facts_list:

                fallback_list = []
                for f in self.fact_bank[:25]:
                    pub = get_clean_publisher(f["source_idx"])
                    fallback_list.append(f"- ({pub}) | {f['fact']}")
                return "\n".join(fallback_list)

            return "\n".join(facts_list[:25])

        accumulated_chapters_context = ""

        for ch_key, ch_spec in self.chapters_outline.items():
            ch_title = ch_spec.get("title", ch_key)
            ch_obj = ch_spec.get("objective", "Overview")

            self.partial_drafts[ch_title] = (
                f"*Status: Actively drafting chapter content in real-time...*\n\n**Objective:** {ch_obj}"
            )

            ch_facts = filter_relevant_facts(ch_title + " " + ch_obj)

            ch_prompt = (
                f"Write '{ch_title}'.\n"
                f"Objective focus: {ch_obj}\n\n"
                f"Here are the previously drafted chapters for context and continuity:\n"
                f"=== PRECEDING CHAPTERS ===\n{accumulated_chapters_context if accumulated_chapters_context else 'None'}\n\n"
                f"Base your writeup strictly on this relevant fact database. Every fact is prefix-wrapped with its "
                f"respective citation key (e.g. `(Sustainableatlas) | ...`). You MUST extract and use these exact "
                f"parenthetical inline keys in your text as your parenthetical inline citations:\n"
                f"{ch_facts}\n\n"
                f"Do not write introductory preambles, conversational chit-chat, or closing remarks. "
                f"Output only the raw chapter content with paragraphs and sub-headings."
            )

            ch_text = ""
            try:
                response = await bot.chat_handler.client.aio.models.generate_content(
                    model=bot.chat_handler.premium_model,
                    contents=ch_prompt,
                    config=types.GenerateContentConfig(
                        system_instruction=system_instruction, temperature=0.4
                    ),
                )
                ch_text = (response.text or "").strip()
            except Exception as e:
                ch_text = f"### {ch_title}\nFailed to draft chapter: {e}"

            self.partial_drafts[ch_title] = ch_text
            accumulated_chapters_context += f"### {ch_title}\n{ch_text}\n\n"

        if self.format_type == "mla":

            current_date_str = datetime.now(timezone.utc).strftime("%d %B %Y")
            complete_markdown = (
                f"Research Analyst\n"
                f"Deep Research Intelligence Assistant\n"
                f"Global Strategic Intelligence\n"
                f"Date: {current_date_str}\n\n"
                f"# {self.clean_title}\n"
                f"----------------------------------------\n\n"
            )
            for ch_title, ch_text in self.partial_drafts.items():
                complete_markdown += f"## {ch_title}\n{ch_text}\n\n"

            if self.works_cited_list:
                complete_markdown += "## Works Cited\n"
                complete_markdown += "----------------------------------------\n\n"
                for citation in self.works_cited_list:
                    complete_markdown += f"{citation}\n\n"
        else:
            sources_str = "\n".join(
                [
                    f"- [Source {i+1}] Title: \"{s['title']}\" | URL: {s['url']}"
                    for i, s in enumerate(self.sources)
                ]
            )
            complete_markdown = f"## Analytical Research Dossier: {self.clean_title}\n"
            complete_markdown += "----------------------------------------\n\n"
            for ch_title, ch_text in self.partial_drafts.items():
                complete_markdown += f"### {ch_title}\n{ch_text}\n\n"
            complete_markdown += f"### 📚 Sources Bibliography\n{sources_str}"

        return complete_markdown
